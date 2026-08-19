from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\working\yuta\yuta-resto")
OUT_DIR = ROOT / "artifacts" / "investor"
OUT_PATH = OUT_DIR / "YUTA_investor_infrastructure_cost_report_2026-08.docx"

NAVY = "153B32"
GREEN = "27866C"
MINT = "EAF5F1"
INK = "17201D"
MUTED = "58645F"
LIGHT = "F3F6F5"
BORDER = "CDD8D4"
WHITE = "FFFFFF"
GOLD = "A87923"
RED = "A13B34"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border_bottom(paragraph, color=GREEN, size="18", space="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instr, separate, text, end])
    paragraph._p.append(field_run)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_body(doc, text: str, bold_lead: str | None = None, after=6) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)


def add_callout(doc, label: str, text: str, fill=MINT, color=NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=fill, size="0")
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run(r, size=11, bold=True, color=color)
    r = p.add_run(text)
    set_run(r, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True


def fill_table(table, rows: list[list[str]], header=True, numeric_cols: set[int] | None = None) -> None:
    numeric_cols = numeric_cols or set()
    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run(run, size=9.2, bold=(header and r_idx == 0), color=(WHITE if header and r_idx == 0 else INK))
            if header and r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)
    if header:
        set_repeat_table_header(table.rows[0])
    set_table_borders(table)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, NAVY),
        ("Heading 2", 13, 12, 6, GREEN),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def setup_section(section, first_page=False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first_page


def add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("YUTA  |  Dự toán chi phí vận hành")
    set_run(r, size=8.5, bold=True, color=MUTED)
    set_paragraph_border_bottom(p, color=BORDER, size="6", space="4")
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    setup_section(doc.sections[0], first_page=True)
    add_running_furniture(doc.sections[0])

    # Editorial cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO NHÀ ĐẦU TƯ")
    set_run(r, size=10, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("YUTA")
    set_run(r, size=34, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Dự toán chi phí vận hành sản phẩm\nở quy mô 40, 80 và 100 nhà hàng")
    set_run(r, size=18, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Bản đánh giá sơ bộ cho giai đoạn hiện tại")
    set_run(r, size=12, italic=True, color=MUTED)

    summary = doc.add_table(rows=1, cols=3)
    set_table_geometry(summary, [3120, 3120, 3120], indent_dxa=120)
    set_table_borders(summary, color=MINT, size="0")
    for idx, (scale, value, unit) in enumerate(
        (("40 NHÀ HÀNG", "€250–500", "/ tháng"), ("80 NHÀ HÀNG", "€400–800", "/ tháng"), ("100 NHÀ HÀNG", "€500–1.000", "/ tháng"))
    ):
        cell = summary.cell(0, idx)
        set_cell_shading(cell, MINT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(scale)
        set_run(r, size=8.5, bold=True, color=GREEN)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value)
        set_run(r, size=17, bold=True, color=NAVY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(unit)
        set_run(r, size=8.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Ngày lập: 18 tháng 8 năm 2026")
    set_run(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Đơn vị chuẩn bị: YUTA")
    set_run(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu làm việc bảo mật - Dự toán chưa kiểm toán")
    set_run(r, size=9, bold=True, color=RED)

    doc.add_page_break()

    add_heading(doc, "1. Tóm tắt điều hành", 1)
    add_callout(
        doc,
        "KẾT LUẬN",
        "Ở quy mô đến 100 nhà hàng, chi phí hạ tầng cloud không phải yếu tố hạn chế chính. Ngân sách hợp lý cho production giai đoạn đầu là khoảng €250–1.000 mỗi tháng, tương đương khoảng €5–12,5 cho mỗi nhà hàng, chưa gồm SMS, thiết bị tại chỗ, nhân sự và hỗ trợ khách hàng.",
    )
    add_body(
        doc,
        "YUTA nên tiếp tục theo mô hình kết hợp: sử dụng dịch vụ được quản lý sẵn cho database, hosting, lưu file, email, AI và giám sát; đồng thời tự sở hữu logic nghiệp vụ, quyền truy cập, dữ liệu theo organization/establishment, quy trình xác nhận và lịch sử hành động.",
    )
    add_body(
        doc,
        "Cách tiếp cận này giữ chi phí cố định thấp, rút ngắn thời gian ra thị trường và tránh việc đội ngũ phải tự vận hành những hạ tầng không tạo lợi thế cạnh tranh. Khi số lượng nhà hàng tăng, chi phí trên mỗi nhà hàng giảm nhờ phần lớn nền tảng được dùng chung.",
    )

    add_heading(doc, "2. Trạng thái sản phẩm được dùng làm cơ sở", 1)
    add_body(doc, "Nền tảng hiện tại. ", bold_lead="Nền tảng hiện tại. ")
    p = doc.paragraphs[-1]
    p.add_run("YUTA là monorepo mô-đun, tách rõ các dịch vụ cloud, Backoffice, các web app công khai và các thành phần vận hành tại nhà hàng. Database cloud và dữ liệu vận hành tại chỗ có ranh giới riêng.")
    for r in p.runs[2:]:
        set_run(r)
    add_body(doc, "Khả năng đã có nền tảng dữ liệu thật. ", bold_lead="Khả năng đã có nền tảng dữ liệu thật. ")
    p = doc.paragraphs[-1]
    p.add_run("Authentication, organization/establishment, quản lý membership, đặt bàn, phản hồi khách hàng và một phần Backoffice đã có đường dữ liệu thực tế.")
    for r in p.runs[2:]:
        set_run(r)
    add_body(doc, "Khả năng còn bị giới hạn. ", bold_lead="Khả năng còn bị giới hạn. ")
    p = doc.paragraphs[-1]
    p.add_run("Một số bề mặt vẫn là prototype hoặc chỉ được phép chạy local. Đặc biệt, tài liệu nhân viên, sổ nhân sự và phân tích hợp đồng bằng AI chưa được xem là production-ready.")
    for r in p.runs[2:]:
        set_run(r)
    add_body(doc, "Ý nghĩa đối với dự toán. ", bold_lead="Ý nghĩa đối với dự toán. ")
    p = doc.paragraphs[-1]
    p.add_run("Con số trong báo cáo là ngân sách kỹ thuật để vận hành nền tảng khi các cổng production cần thiết được hoàn tất; đây không phải xác nhận rằng toàn bộ chức năng hiện đã sẵn sàng bán thương mại.")
    for r in p.runs[2:]:
        set_run(r)

    add_heading(doc, "3. Phạm vi và giả định tính toán", 1)
    assumptions = [
        "Mỗi nhà hàng có khoảng 5 tài khoản Backoffice đang hoạt động.",
        "Mỗi nhà hàng phát sinh khoảng 500 lượt đặt bàn hoặc email giao dịch mỗi tháng.",
        "Mỗi nhà hàng có khoảng 100 review được phân tích và 40 bản nháp trả lời do AI hỗ trợ mỗi tháng.",
        "Mỗi nhà hàng phân tích khoảng 5 hợp đồng nhân viên mỗi tháng; người có quyền vẫn phải duyệt trước khi cập nhật dữ liệu.",
        "Mỗi nhà hàng tạo thêm khoảng 100–200 MB file mỗi tháng.",
        "Không bao gồm SMS, phần cứng, chi phí thanh toán, lương nhân sự, hỗ trợ khách hàng, VAT hoặc yêu cầu Enterprise riêng.",
    ]
    for item in assumptions:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(item)
        set_run(r)

    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    r = p.add_run("Lưu ý: ")
    set_run(r, size=9.5, bold=True, color=GOLD)
    r = p.add_run("Đây là mô hình ngân sách, không phải báo giá nhà cung cấp. Giá nguồn có thể được niêm yết bằng USD; các khoảng EUR đã có biên dự phòng thay vì dùng một tỷ giá cố định.")
    set_run(r, size=9.5, italic=True, color=MUTED)

    add_heading(doc, "4. Dự toán tổng thể", 1)
    total_rows = [
        ["Quy mô", "Ngân sách/tháng", "Chi phí/nhà hàng", "Ngân sách/năm"],
        ["40 nhà hàng", "€250–500", "€6–12,5", "€3.000–6.000"],
        ["80 nhà hàng", "€400–800", "€5–10", "€4.800–9.600"],
        ["100 nhà hàng", "€500–1.000", "€5–10", "€6.000–12.000"],
    ]
    table = doc.add_table(rows=len(total_rows), cols=4)
    set_table_geometry(table, [1900, 2450, 2450, 2560])
    fill_table(table, total_rows, numeric_cols={1, 2, 3})
    add_body(doc, "Khoảng trên đã bao gồm một phần dự phòng tăng tải và dao động hóa đơn. Mức thực tế sẽ được điều chỉnh sau khi có dữ liệu sử dụng production trong ba tháng đầu.", after=2)

    add_heading(doc, "5. Cấu trúc chi phí hàng tháng", 1)
    detail_rows = [
        ["Hạng mục", "40 nhà hàng", "80 nhà hàng", "100 nhà hàng"],
        ["Hosting các web app", "€60–120", "€90–180", "€110–220"],
        ["Neon PostgreSQL", "€35–80", "€60–140", "€80–180"],
        ["Email giao dịch", "€30–45", "€60–80", "€75–100"],
        ["Lưu file EU và quét virus", "€5–20", "€10–30", "€15–40"],
        ["Giám sát lỗi, log, cảnh báo", "€25–60", "€30–80", "€40–100"],
        ["AI: review và hợp đồng", "€5–20", "€10–40", "€15–50"],
        ["Backup và chi phí phụ", "€20–50", "€30–70", "€40–90"],
        ["Dự phòng tăng tải", "€50–105", "€110–180", "€125–220"],
        ["Tổng dự kiến", "€250–500", "€400–800", "€500–1.000"],
    ]
    table = doc.add_table(rows=len(detail_rows), cols=4)
    set_table_geometry(table, [3300, 2020, 2020, 2020])
    fill_table(table, detail_rows, numeric_cols={1, 2, 3})
    for cell in table.rows[-1].cells:
        set_cell_shading(cell, MINT)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9.4, bold=True, color=NAVY)

    add_heading(doc, "6. Diễn giải kinh tế đơn vị", 1)
    add_body(doc, "Chi phí cloud trên mỗi nhà hàng giảm từ khoảng €6–12,5 ở mốc 40 nhà hàng xuống khoảng €5–10 ở mốc 80–100 nhà hàng. Điều này phản ánh khả năng dùng chung hosting, database, monitoring và các service trung tâm.")
    add_body(doc, "Ở mốc 100 nhà hàng, doanh thu chỉ cần cao hơn €10 mỗi nhà hàng mỗi tháng là đã vượt chi phí hạ tầng kỹ thuật trực tiếp trong mô hình cơ sở. Tuy nhiên, đây không phải biên lợi nhuận: chi phí hỗ trợ, bán hàng, onboarding, phần cứng, pháp lý và nhân sự vẫn phải được tính riêng.")
    add_callout(
        doc,
        "ĐIỂM QUAN TRỌNG CHO NHÀ ĐẦU TƯ",
        "Rủi ro tài chính chính không nằm ở token AI hay dung lượng file. Các yếu tố dễ làm tăng chi phí hơn là SMS, hỗ trợ vận hành, yêu cầu Enterprise và phần cứng tại nhà hàng.",
    )

    add_heading(doc, "7. Các khoản có thể làm thay đổi đáng kể ngân sách", 1)
    add_heading(doc, "7.1 SMS nhắc đặt bàn", 2)
    add_body(doc, "Nếu mỗi nhà hàng gửi 500 SMS mỗi tháng và chi phí giả định là €0,05–0,10 cho mỗi tin, SMS sẽ bổ sung khoảng €1.000–2.000/tháng ở 40 nhà hàng, €2.000–4.000 ở 80 nhà hàng và €2.500–5.000 ở 100 nhà hàng.")
    add_body(doc, "Vì vậy SMS nên được bán như add-on hoặc theo quota, không nên để sử dụng không giới hạn trong gói cơ bản.")

    add_heading(doc, "7.2 Production tăng cường", 2)
    enterprise_rows = [
        ["Quy mô", "Ngân sách production tăng cường"],
        ["40 nhà hàng", "€900–1.300/tháng"],
        ["80 nhà hàng", "€1.050–1.550/tháng"],
        ["100 nhà hàng", "€1.150–1.800/tháng"],
    ]
    table = doc.add_table(rows=len(enterprise_rows), cols=2)
    set_table_geometry(table, [3600, 5760])
    fill_table(table, enterprise_rows, numeric_cols={1})
    add_body(doc, "Mức này phù hợp khi hợp đồng khách hàng yêu cầu SLA, private networking, thời gian khôi phục dài, hỗ trợ ưu tiên hoặc kiểm soát bảo mật nâng cao. YUTA chưa cần mua toàn bộ mức Enterprise trước khi nhu cầu thương mại thực tế xuất hiện.")

    add_heading(doc, "7.3 Các khoản nằm ngoài mô hình", 2)
    excluded = [
        "Thiết bị tại nhà hàng, máy in, màn hình, UPS, router, lắp đặt và bảo trì tại chỗ.",
        "Nhân sự kỹ thuật, chăm sóc khách hàng, trực sự cố, bán hàng, kế toán, pháp lý và DPO.",
        "Chi phí tạo nhiều hình ảnh AI chất lượng cao hoặc các chiến dịch marketing lớn.",
        "Chi phí thay đổi do tỷ giá, thuế, vùng lưu trữ hoặc hợp đồng nhà cung cấp.",
    ]
    for item in excluded:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)

    add_heading(doc, "8. Chiến lược hạ tầng đề xuất", 1)
    strategy_rows = [
        ["Thành phần", "Lựa chọn hiện tại", "Phần YUTA phải sở hữu"],
        ["Database", "Neon PostgreSQL được quản lý", "Schema, tenancy, authorization, backup test"],
        ["File", "Kho riêng tư tại EU", "Metadata, quyền truy cập, retention, adapter thay thế"],
        ["AI/OCR", "Một nhà cung cấp AI bên ngoài", "Prompt, dữ liệu gửi đi, validation, review, audit"],
        ["Email", "Nhà cung cấp email giao dịch", "Template, trạng thái gửi, retry, quyền dữ liệu"],
        ["Giám sát", "Dịch vụ monitoring tập trung", "Mức cảnh báo, runbook, người chịu trách nhiệm"],
        ["Logic sản phẩm", "YUTA tự xây", "Nghiệp vụ, phân quyền, lịch sử và trải nghiệm người dùng"],
    ]
    table = doc.add_table(rows=len(strategy_rows), cols=3)
    set_table_geometry(table, [1900, 3000, 4460])
    fill_table(table, strategy_rows)

    add_body(doc, "Mỗi nhà cung cấp cần được đặt sau một service/adapter do YUTA kiểm soát. Việc đổi kho file hoặc AI khi đó chỉ thay adapter và cấu hình, không viết lại nghiệp vụ hoặc giao diện.")

    add_heading(doc, "9. Rủi ro và biện pháp kiểm soát", 1)
    risk_rows = [
        ["Rủi ro", "Tác động", "Biện pháp"],
        ["Chi phí cloud tăng ngoài dự kiến", "Giảm biên lợi nhuận", "Budget cap, cảnh báo, quota và theo dõi theo tenant"],
        ["Phụ thuộc nhà cung cấp", "Khó chuyển đổi", "Adapter riêng, dữ liệu xuất được, không để UI gọi trực tiếp"],
        ["Dữ liệu nhân viên nhạy cảm", "Rủi ro pháp lý và uy tín", "EU storage, quyền OWNER, audit, retention và phê duyệt trước production"],
        ["AI trả kết quả sai", "Sai hồ sơ hoặc phản hồi", "Structured output, validation và người dùng xác nhận"],
        ["Khả năng chưa production-ready", "Cam kết thương mại quá sớm", "Giữ fail-closed và chỉ mở khi đủ bằng chứng vận hành"],
        ["SMS/phần cứng vượt dự toán", "Chi phí theo nhà hàng cao", "Add-on, quota và tách khỏi giá nền tảng"],
    ]
    table = doc.add_table(rows=len(risk_rows), cols=3)
    set_table_geometry(table, [2350, 2550, 4460])
    fill_table(table, risk_rows)

    add_heading(doc, "10. Kế hoạch kiểm chứng 90 ngày", 1)
    steps = [
        "Tháng 1: hoàn tất baseline production, xác định vùng dữ liệu, backup, giám sát và giới hạn chi phí cho từng dịch vụ.",
        "Tháng 2: chạy pilot có kiểm soát, đo database compute, request, email, file, AI và khối lượng hỗ trợ thực tế theo nhà hàng.",
        "Tháng 3: so sánh hóa đơn với mô hình này, cập nhật chi phí đơn vị và quyết định thời điểm cần nâng cấp gói production.",
    ]
    for item in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)

    add_callout(
        doc,
        "TIÊU CHÍ QUYẾT ĐỊNH",
        "Chỉ nâng cấp sang gói Enterprise khi một yêu cầu hợp đồng, bảo mật hoặc độ tin cậy có giá trị lớn hơn phần chi phí tăng thêm; không nâng cấp chỉ vì số lượng nhà hàng đạt một con số tròn.",
    )

    add_heading(doc, "11. Kết luận", 1)
    add_body(doc, "Về mặt chi phí, mô hình hạ tầng đề xuất có thể hỗ trợ quy mô 40–100 nhà hàng với chi phí trung tâm thấp so với chi phí nhân sự và hỗ trợ. Đây là dự toán tài chính, không phải chứng nhận năng lực production. Mô hình dịch vụ được quản lý giúp công ty tập trung vốn và thời gian vào sản phẩm, dữ liệu, trải nghiệm người dùng và vận hành nhà hàng.")
    add_body(doc, "Ngân sách khuyến nghị cho giai đoạn hiện tại là €250–1.000 mỗi tháng, kèm giới hạn chi tiêu và ba tháng đo lường thực tế. SMS, phần cứng, nhân sự và production Enterprise phải được quản lý như các dòng ngân sách riêng.")

    add_heading(doc, "12. Nguồn tham khảo và giới hạn", 1)
    sources = [
        ("Vercel pricing", "https://vercel.com/pricing", "Gói Pro và cách tính phần usage vượt mức."),
        ("Neon pricing", "https://neon.com/pricing", "Compute, storage, restore window và ví dụ tải."),
        ("Postmark pricing", "https://postmarkapp.com/pricing", "Email giao dịch và chi phí vượt quota."),
        ("OpenAI GPT-5.6 Luna", "https://developers.openai.com/api/docs/models/gpt-5.6-luna", "Giá token cho workload có chi phí thấp."),
        ("OpenAI data controls", "https://developers.openai.com/api/docs/guides/your-data", "Retention, data residency và các điều kiện kiểm soát dữ liệu."),
        ("AWS GuardDuty pricing", "https://aws.amazon.com/guardduty/pricing/", "Quét malware cho file tải lên S3."),
        ("AWS S3 encryption", "https://docs.aws.amazon.com/AmazonS3/latest/userguide/bucket-encryption.html", "Mã hóa mặc định và lựa chọn KMS."),
    ]
    for label, url, note in sources:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        add_hyperlink(p, label, url)
        r = p.add_run(f" — {note}")
        set_run(r, size=9.5, color=MUTED)

    add_body(doc, "Tất cả dự toán trong báo cáo là chỉ dẫn lập kế hoạch tại ngày 18/08/2026. Chi phí thực tế phụ thuộc mức sử dụng, vùng triển khai, tỷ giá, thuế, điều khoản hợp đồng và lựa chọn kỹ thuật cuối cùng.", after=2)

    doc.core_properties.title = "YUTA - Dự toán chi phí vận hành sản phẩm"
    doc.core_properties.subject = "Báo cáo sơ bộ dành cho nhà đầu tư"
    doc.core_properties.author = "YUTA"
    doc.core_properties.keywords = "YUTA, investor, infrastructure, operating cost"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
